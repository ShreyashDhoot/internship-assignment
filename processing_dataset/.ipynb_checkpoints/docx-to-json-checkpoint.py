import os
import json
import re
from pathlib import Path
import docx

def parse_docx_content(file_path, folder_artist_id, folder_artist_name, category_name):
    doc = docx.Document(file_path)
    
    # 1. Extract raw text from paragraphs and tables
    raw_lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t: raw_lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in raw_lines: raw_lines.append(t)

    # 2. Split soft returns / embedded newlines (\n) into distinct lines
    lines = []
    for line in raw_lines:
        for sub in line.split('\n'):
            cleaned = sub.strip()
            if cleaned:
                lines.append(cleaned)

    profile = {
        "artist_id": folder_artist_id,
        "name": folder_artist_name,
        "category": category_name,          # folder-based, GROUND TRUTH for filtering -- never overwritten below
        "docx_stated_category": None,        # whatever the docx itself says, kept separately as a claim to verify
        "location": None,
        "work_preference": None,
        "bio": None,
        "portfolio": []
    }

    # Extract Header Line (e.g., "P03 / Leena Thomas")
    if lines and ("/" in lines[0] or "—" in lines[0] or "-" in lines[0]):
        header_parts = re.split(r'[/—\-]', lines[0], maxsplit=1)
        if len(header_parts) == 2:
            extracted_id = header_parts[0].strip()
            extracted_name = header_parts[1].strip()
            if re.match(r'^[A-Z0-9_]+$', extracted_id, re.IGNORECASE):
                profile["artist_id"] = extracted_id
                profile["name"] = extracted_name

    current_section = None

    for i, line in enumerate(lines):
        line_clean = line.strip().lower().rstrip(":")

        if line_clean == "bio":
            current_section = "bio"
            continue
        elif line_clean in ["portfolio", "portfolio:"]:
            current_section = "portfolio"
            continue

        if ":" in line or "Preference-" in line:
            parts = re.split(r'[:\-]', line, maxsplit=1)
            key = parts[0].strip().lower()
            val = parts[1].strip() if len(parts) > 1 else ""

            if "location" in key and val:
                profile["location"] = val
                current_section = None
                continue
            elif "preference" in key and val:
                profile["work_preference"] = val
                current_section = None
                continue
            elif "category" in key and val:
                # IMPORTANT: do NOT overwrite profile["category"] here.
                # That field is the folder-based ground truth (musicians/
                # photographers/video_editors), used for filtering, and
                # is deliberately set once above and left alone -- the
                # docx can state something different or broader (e.g.
                # VO4_Shivam_media's docx says "Visual Artist" while its
                # folder is video_editors) and that disagreement is
                # itself a signal worth keeping, not silently erasing.
                profile["docx_stated_category"] = val
                current_section = None
                continue

        if line_clean == "category" and i + 1 < len(lines):
            # Same reasoning as above: this is the docx's own stated
            # category (label-alone-then-next-line layout, e.g. the
            # musician docx format), never the folder-based ground truth.
            profile["docx_stated_category"] = lines[i+1].strip()
        elif line_clean == "location" and i + 1 < len(lines):
            profile["location"] = lines[i+1].strip()
        elif line_clean in ["work preference", "work preference-onsite"]:
            if i + 1 < len(lines) and not profile["work_preference"]:
                profile["work_preference"] = lines[i+1].strip()

        if current_section == "bio":
            profile["bio"] = f"{profile['bio']} {line}".strip() if profile["bio"] else line
        elif current_section == "portfolio":
            profile["portfolio"].append(line)

    return profile


if __name__ == "__main__":
    # --- RELATIVE PATH CALCULATION ---
    # Gets the script's directory (.../internship-assignment/processing_dataset)
    SCRIPT_DIR = Path(__file__).resolve().parent
    
    # Navigates to the repo root (.../internship-assignment)
    REPO_ROOT = SCRIPT_DIR.parent
    
    # Relative path to artist profiles
    dataset_path = REPO_ROOT / "Data-set" / "artist_profiles"
    
    # Fallback check in case the script is placed directly in the repo root
    if not dataset_path.exists():
        dataset_path = SCRIPT_DIR / "Data-set" / "artist_profiles"
        
    if not dataset_path.exists():
        print(f"[ERROR] Could not locate dataset at: {dataset_path}")
        exit(1)

    # Separate output folder for generated JSON, kept apart from the
    # source docx/media folders rather than writing profile.json inside
    # each artist's own source directory.
    output_dir = REPO_ROOT / "artist_profiles"
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"Target dataset directory: {dataset_path}")
    print(f"Output directory: {output_dir}\n")

    written_filenames = set()

    # Process all directories recursively
    for docx_file in dataset_path.rglob("*.docx"):
        if docx_file.name.startswith("~$"):
            continue
            
        artist_dir = docx_file.parent
        category_dir = artist_dir.parent
        
        category_name = category_dir.name
        folder_name = artist_dir.name
        
        parts = folder_name.split("_", 1)
        folder_artist_id = parts[0]
        folder_artist_name = parts[1].replace("_", " ") if len(parts) > 1 else "Unknown"

        # Run extraction
        profile = parse_docx_content(docx_file, folder_artist_id, folder_artist_name, category_name)

        # Name the output file after the ARTIST FOLDER, not the docx-
        # embedded artist_id -- the folder name is guaranteed unique
        # (it's a real directory on disk), whereas artist_id can come
        # from docx-internal text that parse_docx_content may override
        # (see lines 40-47) and is known to collide across artists in
        # this dataset (two different video editors' docx both say
        # "V03"). Using the folder name as the filename avoids one
        # artist's profile.json silently overwriting another's.
        safe_filename = "".join(
            c if c.isalnum() or c in ("_", "-") else "_" for c in folder_name
        )
        json_filename = f"{safe_filename}.json"

        if json_filename in written_filenames:
            print(f"[WARNING] Filename collision: {json_filename} (from folder '{folder_name}') "
                  f"already written this run -- check for duplicate/near-duplicate folder names.")
        written_filenames.add(json_filename)

        json_path = output_dir / json_filename
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(profile, f, indent=4)
            
        print(f"Created JSON: {profile['artist_id']} - {profile['name']} -> {json_path.name}")

    print(f"\nExtraction complete. {len(written_filenames)} profile JSON files generated in {output_dir}")