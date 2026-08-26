import os
import json
import re
import sqlite3
from pathlib import Path
import docx


def parse_docx_content(file_path, folder_artist_id, folder_artist_name, category_name):
    doc = docx.Document(file_path)
    
    # 1. Extract raw text from paragraphs and tables
    raw_lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t: raw_lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in raw_lines: raw_lines.append(t)

    # 2. Split soft returns / embedded newlines (\n) into distinct lines
    lines = []
    for line in raw_lines:
        for sub in line.split('\n'):
            cleaned = sub.strip()
            if cleaned:
                lines.append(cleaned)

    profile = {
        "artist_id": folder_artist_id,
        "name": folder_artist_name,
        "category": category_name,
        "location": None,
        "work_preference": None,
        "bio": None,
        "portfolio": []
    }

    # Extract Header Line (e.g., "P03 / Leena Thomas")
    if lines and ("/" in lines[0] or "—" in lines[0] or "-" in lines[0]):
        header_parts = re.split(r'[/—\-]', lines[0], maxsplit=1)
        if len(header_parts) == 2:
            extracted_id = header_parts[0].strip()
            extracted_name = header_parts[1].strip()
            if re.match(r'^[A-Z0-9_]+$', extracted_id, re.IGNORECASE):
                profile["artist_id"] = extracted_id
                profile["name"] = extracted_name

    current_section = None

    for i, line in enumerate(lines):
        # Clean the line for section checking (strips colons like "Bio:" -> "bio")
        line_clean = line.strip().lower().rstrip(":")

        # --- STEP 1: Check section headers BEFORE splitting colons ---
        if line_clean == "bio":
            current_section = "bio"
            continue
        elif line_clean in ["portfolio", "portfolio:"]:
            current_section = "portfolio"
            continue

        # --- STEP 2: Handle Inline Key-Value Pairs ---
        if ":" in line or "Preference-" in line:
            parts = re.split(r'[:\-]', line, maxsplit=1)
            key = parts[0].strip().lower()
            val = parts[1].strip() if len(parts) > 1 else ""

            if "location" in key and val:
                profile["location"] = val
                current_section = None
                continue
            elif "preference" in key and val:
                profile["work_preference"] = val
                current_section = None
                continue
            elif "category" in key and val:
                profile["category"] = val
                current_section = None
                continue

        # --- STEP 3: Handle Multiline Headers (Key on line i, Value on line i+1) ---
        if line_clean == "category" and i + 1 < len(lines):
            profile["category"] = lines[i+1].strip()
        elif line_clean == "location" and i + 1 < len(lines):
            profile["location"] = lines[i+1].strip()
        elif line_clean in ["work preference", "work preference-onsite"]:
            if i + 1 < len(lines) and not profile["work_preference"]:
                profile["work_preference"] = lines[i+1].strip()

        # --- STEP 4: Accumulate multiline content (Bio or Portfolio) ---
        if current_section == "bio":
            profile["bio"] = f"{profile['bio']} {line}".strip() if profile["bio"] else line
        elif current_section == "portfolio":
            profile["portfolio"].append(line)

    return profile
    
def setup_sqlite_db(db_path="artists.db"):
    """Creates a local SQLite database for native JSON queries."""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS artist_documents (
            artist_id TEXT PRIMARY KEY,
            data TEXT CHECK(json_valid(data))
        )
    ''')
    conn.commit()
    return conn


def save_record(conn, artist_dir, profile):
    """Saves both a local profile.json and inserts the JSON string into SQLite."""
    # 1. Save JSON file inside artist directory
    json_path = artist_dir / "profile.json"
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(profile, f, indent=4)
        
    # 2. Insert into SQLite table
    cursor = conn.cursor()
    cursor.execute('''
        INSERT OR REPLACE INTO artist_documents (artist_id, data)
        VALUES (?, ?)
    ''', (profile["artist_id"], json.dumps(profile)))
    conn.commit()


# MAIN EXECUTION BLOCK
if __name__ == "__main__":
    dataset_path = Path(r"D:\INTERNSHIP-assigment\Data-set\artist_profiles")
    db_conn = setup_sqlite_db("artists.db")
    
    # Process all directories recursively
    for docx_file in dataset_path.rglob("*.docx"):
        # Skip temporary Office files
        if docx_file.name.startswith("~$"):
            continue
            
        artist_dir = docx_file.parent
        category_dir = artist_dir.parent
        
        # Derive initial metadata from folder structure
        category_name = category_dir.name
        folder_name = artist_dir.name
        
        parts = folder_name.split("_", 1)
        folder_artist_id = parts[0]
        folder_artist_name = parts[1].replace("_", " ") if len(parts) > 1 else "Unknown"

        # Run extraction
        profile = parse_docx_content(docx_file, folder_artist_id, folder_artist_name, category_name)
        
        # Save output to disk & database
        save_record(db_conn, artist_dir, profile)
        print(f"Processed: {profile['artist_id']} - {profile['name']}")

    db_conn.close()
    print("\nExtraction complete. Saved all records to JSON files and artists.db.")